import os
from datetime import datetime

import pandas as pd
import docx
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ── Paths ─────────────────────────────────────────────────────────────────────
TEMPLATE_PATH = r"z:\HTOC\HTOC Reports\I&W Reports\5. I&W Staging\I&W Report Template.docx"
OUTPUT_DIR    = r"z:\HTOC\HTOC Reports\I&W Reports\5. I&W Staging\Generated Reports"
os.makedirs(OUTPUT_DIR, exist_ok=True)

# ── Internal helpers (new, private) ───────────────────────────────────────────
def _to_date_str(val):
    """Return YYYY-MM-DD or 'N/A'."""
    if pd.isna(val) or val == "N/A":
        return "N/A"
    try:
        ts = pd.to_datetime(val, errors="coerce", utc=True)
        if pd.isna(ts):
            return "N/A"
        return ts.date().isoformat()
    except Exception:
        return "N/A"

def _replace_placeholder_with_text(paragraph, placeholder, replacement):
    """
    Replace {{placeholder}} inside a paragraph without nuking run formatting.
    Returns True if replaced in this paragraph.
    """
    if placeholder not in paragraph.text:
        return False
    # rebuild, replace, then write back as a single run (safe + simple)
    full = "".join(run.text for run in paragraph.runs)
    full = full.replace(placeholder, str(replacement))
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = full
    else:
        paragraph.add_run(full)
    return True

def _replace_placeholder_with_hyperlinks(paragraph, placeholder, links):
    """
    Replace {{placeholder}} with list of hyperlinks (one per line).
    """
    if placeholder not in paragraph.text:
        return False
    # clear content, remove placeholder
    for r in paragraph.runs:
        r.text = ""
    paragraph.text = paragraph.text.replace(placeholder, "")
    links = [l for l in (links or []) if isinstance(l, str) and l.strip()]
    if not links:
        paragraph.add_run("N/A")
        return True
    for i, url in enumerate(links):
        add_hyperlink(paragraph, url)
        if i < len(links) - 1:
            paragraph.add_run().add_break()
    return True

def _find_indicator_table(doc: Document):
    """Find the table whose first row contains 'Indicators/Identifiers'."""
    for tbl in doc.tables:
        try:
            header = " | ".join(c.text for c in tbl.rows[0].cells)
        except IndexError:
            continue
        if "Indicators/Identifiers" in header:
            return tbl
    return None

def _normalize_group_id(series: pd.Series) -> pd.Series:
    """Turn '', 'None', 'null', 'nan', whitespace into <NA>; strip strings."""
    if series is None or not isinstance(series, pd.Series):
        return pd.Series([], dtype="string")
    out = series.astype("string").str.strip()
    out = out.mask(out.isna() | out.eq("") | out.str.lower().isin({"none", "null", "nan"}))
    return out

# ── Utilities (names kept) ────────────────────────────────────────────────────
def consolidate_sources(vt_df: pd.DataFrame | None, otx_df: pd.DataFrame | None) -> pd.DataFrame:
    """Consolidate links from VT/OTX per search_term -> 'sources' (comma-joined)."""
    def _links_from(df, out_col):
        if df is None or not isinstance(df, pd.DataFrame) or df.empty:
            return pd.DataFrame(columns=["search_term", out_col])
        if "search_term" not in df.columns or "link" not in df.columns:
            return pd.DataFrame(columns=["search_term", out_col])
        out = (
            df.groupby("search_term")["link"]
              .apply(lambda x: ", ".join(x.dropna().astype(str)))
              .reset_index()
              .rename(columns={"link": out_col})
        )
        return out

    vt_links  = _links_from(vt_df, "vt_links")
    otx_links = _links_from(otx_df, "otx_links")

    if vt_links.empty and otx_links.empty:
        return pd.DataFrame(columns=["search_term", "sources"])

    consolidated = pd.merge(vt_links, otx_links, on="search_term", how="outer")
    consolidated["vt_links"]  = consolidated["vt_links"].fillna("")
    consolidated["otx_links"] = consolidated["otx_links"].fillna("")
    consolidated["sources"] = consolidated[["vt_links", "otx_links"]].apply(
        lambda x: ", ".join([p for p in x if p]).strip(", ").strip(), axis=1
    )
    return consolidated[["search_term", "sources"]]

def extract_date(ts) -> str:
    """Return YYYY-MM-DD from a datetime/str; 'N/A' if parse fails."""
    if ts is None or (isinstance(ts, float) and pd.isna(ts)) or ts == "N/A":
        return "N/A"
    try:
        dt = pd.to_datetime(ts, errors="coerce", utc=True)
        if pd.isna(dt):
            return "N/A"
        return dt.date().isoformat()
    except Exception:
        return "N/A"

def add_hyperlink(paragraph, url: str):
    """Add a clickable hyperlink run to a paragraph styled as a hyperlink."""
    if not url:
        return
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    new_run.append(rPr)

    t = OxmlElement("w:t")
    t.text = url
    new_run.append(t)
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

# ── Word population (names kept) ──────────────────────────────────────────────
def populate_table(table, data: pd.DataFrame):
    """
    Fill the IOCs table. Expects columns:
      - search_term, type, observed_date, observed_by_otx|observed_by, notes
    """
    for _, row in data.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(row.get("search_term", "N/A"))
        cells[1].text = str(row.get("type", "N/A"))
        cells[2].text = _to_date_str(row.get("observed_date", "N/A"))

        # Build observed_by stacked (one per line)
        observed_by_list = []
        if "observed_by_otx" in row and pd.notna(row["observed_by_otx"]):
            observed_by_list += [v.strip() for v in str(row["observed_by_otx"]).split(",") if v.strip()]
        if "observed_by" in row and pd.notna(row["observed_by"]):
            observed_by_list += [v.strip() for v in str(row["observed_by"]).split(",") if v.strip()]
        observed_by = "\n".join(sorted(set(observed_by_list))) if observed_by_list else "N/A"
        cells[3].text = observed_by

        cells[4].text = str(row.get("notes", "") if pd.notna(row.get("notes", "")) else "")

def fill_word_template(template_path: str,
                       output_path: str,
                       df: pd.DataFrame,
                       recent_tags: pd.DataFrame | None = None):
    """
    Fill template with df rows and replace placeholders outside the table.
    Placeholders:
      {{indicator}}, {{asn}}, {{whois}}, {{partners}}, {{weblink}}, {{sources}}
    """
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template not found: {template_path}")

    doc = Document(template_path)

    # Find the target table by its header text
    target_table = _find_indicator_table(doc)
    if target_table is not None and isinstance(df, pd.DataFrame) and not df.empty:
        populate_table(target_table, df)

    # Compute single-value fields (prefer first non-null across rows)
    def _first_non_null(col):
        return str(df[col].dropna().astype(str).iloc[0]) if (col in df.columns and df[col].notna().any()) else "N/A"

    # Indicators in this section (may be many if you later group)
    indicators = list(pd.Series(df["search_term"].dropna().astype(str)).unique()) if "search_term" in df.columns else ["N/A"]
    search_term_first = indicators[0]
    indicator_count   = len(indicators)
    ip_placeholder_text = search_term_first if indicator_count == 1 else f"{search_term_first} (+{indicator_count - 1} more)"

    asn_value   = _first_non_null("asn")
    whois_value = _first_non_null("whois")

    # Partners (from recent_tags by summary/search_term)
    partners_value = "N/A"
    if isinstance(recent_tags, pd.DataFrame) and not recent_tags.empty:
        st = search_term_first
        if "summary" in recent_tags.columns:
            partners_row = recent_tags[recent_tags["summary"] == st]
        else:
            partners_row = recent_tags[recent_tags.get("search_term", pd.Series([], dtype=object)) == st]
        if not partners_row.empty and "partners" in partners_row.columns:
            pv = partners_row["partners"].dropna().astype(str).iloc[0]
            partners_value = pv if pv else "N/A"

    # Pick a single weblink if available
    weblink_value = ""
    for col in ("webLink", "link"):
        if col in df.columns and df[col].notna().any():
            weblink_value = str(df[col].dropna().astype(str).iloc[0])
            if weblink_value:
                break

    # Collect unique sources (flatten comma-joined cells)
    sources_flat = []
    if "sources" in df.columns and not df["sources"].isna().all():
        for srcs in df["sources"].dropna().astype(str).unique():
            for src in [s.strip() for s in srcs.split(",") if s.strip()]:
                sources_flat.append(src)
    sources_flat = list(dict.fromkeys(sources_flat))  # de-dup preserve order

    # Replace placeholders (safe)
    for para in doc.paragraphs:
        _replace_placeholder_with_text(para, "{{indicator}}", ip_placeholder_text)
        _replace_placeholder_with_text(para, "{{asn}}",       asn_value)
        _replace_placeholder_with_text(para, "{{whois}}",     whois_value)
        _replace_placeholder_with_text(para, "{{partners}}",  partners_value)

        if "{{weblink}}" in para.text:
            for r in para.runs:
                r.text = ""
            para.text = para.text.replace("{{weblink}}", "")
            if weblink_value:
                add_hyperlink(para, weblink_value)
            else:
                para.add_run("N/A")

        _replace_placeholder_with_hyperlinks(para, "{{sources}}", sources_flat)

    # Save to the provided output_path
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"Saved report: {output_path}")

# ── Orchestration (name kept) ────────────────────────────────────────────────
def generate_report(vt_df: pd.DataFrame | None,
                    otx_df: pd.DataFrame | None,
                    recent_tags: pd.DataFrame | None):
    """
    Normalize, merge, and emit reports. If group_id is present, emit one doc per group_id;
    otherwise emit one doc per search_term.
    """
    # Normalize VT
    if isinstance(vt_df, pd.DataFrame) and not vt_df.empty:
        vt_df = vt_df.rename(columns={"ip": "search_term", "webLink": "link"})
    else:
        vt_df = None
    
    # Normalize recent_tags
    if isinstance(recent_tags, pd.DataFrame) and not recent_tags.empty:
        recent_norm = recent_tags.copy()
        if "summary" in recent_norm.columns:
            recent_norm = recent_norm.rename(columns={"summary": "search_term"})
        keep_cols = [c for c in ["search_term", "observations", "type", "partners", "observed_date", "webLink", "group_id"]
                     if c in recent_norm.columns]
        recent_norm = recent_norm[keep_cols]
    else:
        recent_norm = pd.DataFrame(columns=["search_term"])

    print("recent_tags DataFrame (first 20 rows):")
    print(recent_tags.head(20))
    # Combine VT & OTX
    if vt_df is None or vt_df.empty:
        combined_df = otx_df.copy() if isinstance(otx_df, pd.DataFrame) else pd.DataFrame(columns=["search_term"])
    else:
        if isinstance(otx_df, pd.DataFrame):
            combined_df = pd.merge(vt_df, otx_df, on="search_term", how="outer", suffixes=("_vt", "_otx"))
        else:
            combined_df = vt_df.copy()

    # Consolidated sources
    sources_df = consolidate_sources(vt_df, otx_df)
    combined_df = pd.merge(combined_df, sources_df, on="search_term", how="left")

    # Merge recent tags (brings in group_id if present there)
    combined_df = pd.merge(combined_df, recent_norm, on="search_term", how="left")

    # Also map group_id from the original recent_tags by summary -> search_term (guarded)
    if (
        isinstance(recent_tags, pd.DataFrame) and not recent_tags.empty and
        "summary" in recent_tags.columns and "group_id" in recent_tags.columns and
        "search_term" in combined_df.columns
    ):
        group_id_map = recent_tags.set_index("summary")["group_id"]
        combined_df["group_id"] = combined_df["search_term"].map(group_id_map).fillna(combined_df.get("group_id"))

    # ── Coalesce 'type' from possible columns (prevents N/A in table) ─────────
    type_candidates = [c for c in ["type", "type_vt", "type_otx", "indicatorType", "indicator_type"]
                       if c in combined_df.columns]
    if type_candidates:
        combined_df["type"] = (
            combined_df[type_candidates]
            .astype(object)
            .bfill(axis=1)
            .iloc[:, 0]
        )
    else:
        combined_df["type"] = "N/A"

    # ── Normalize group_id so grouping works even with odd strings ────────────
    if "group_id" in combined_df.columns:
        combined_df["group_id"] = _normalize_group_id(combined_df["group_id"])

    # Debug: Display combined_df for inspection
    print("Combined DataFrame (first 20 rows):")
    print(combined_df.head(20))

    # Output folder by date
    date_folder = datetime.now().strftime("%Y-%m-%d")
    folder_path = os.path.join(OUTPUT_DIR, date_folder)
    os.makedirs(folder_path, exist_ok=True)

    # ── Grouping logic ────────────────────────────────────────────────────────
    has_group = "group_id" in combined_df.columns and combined_df["group_id"].notna().any()

    if has_group:
        # 1) One report per non-null group_id
        for gid, gdf in combined_df.dropna(subset=["group_id"]).groupby("group_id", dropna=True):
            # Sort for nice, stable output (optional)
            gdf_sorted = gdf.sort_values(["search_term", "observed_date"], na_position="last")
            base_name = f"Group_{gid}"
            sanitized = str(base_name).replace(":", "_").replace("/", "_")
            output_file = os.path.join(folder_path, f"I&W_Report_{sanitized}.docx")
            fill_word_template(TEMPLATE_PATH, output_file, gdf_sorted, recent_tags=recent_tags)

        # 2) Orphans (no group_id) → one doc per indicator
        orphan_df = combined_df[combined_df["group_id"].isna()] if "group_id" in combined_df.columns else combined_df.iloc[0:0]
        if "search_term" in orphan_df.columns and orphan_df["search_term"].notna().any():
            for indicator in orphan_df["search_term"].dropna().unique():
                idf = orphan_df[orphan_df["search_term"] == indicator].copy()
                sanitized = str(indicator).replace(":", "_").replace("/", "_")
                output_file = os.path.join(folder_path, f"I&W_Report_{sanitized}.docx")
                fill_word_template(TEMPLATE_PATH, output_file, idf, recent_tags=recent_tags)
    else:
        # Fallback: one report per indicator
        indicators = list(combined_df["search_term"].dropna().unique()) if "search_term" in combined_df.columns else []
        if not indicators:
            print(" No indicators found. No reports were generated.")
            return
        for indicator in indicators:
            idf = combined_df[combined_df["search_term"] == indicator].copy()
            sanitized = str(indicator).replace(":", "_").replace("/", "_")
            output_file = os.path.join(folder_path, f"I&W_Report_{sanitized}.docx")
            fill_word_template(TEMPLATE_PATH, output_file, idf, recent_tags=recent_tags)

# ── Entry point (names kept) ─────────────────────────────────────────────────
if __name__ == "__main__":
    # Expect vt_df, otx_df, recent_tags to be defined earlier in your pipeline.
    vt_df = globals().get("vt_df")
    otx_df = globals().get("otx_df")
    recent_tags = globals().get("recent_tags")
    
    generate_report(vt_df, otx_df, recent_tags)
