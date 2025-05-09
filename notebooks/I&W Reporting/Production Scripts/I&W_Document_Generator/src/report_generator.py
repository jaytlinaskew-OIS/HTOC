# src/report_generator.py

from docx import Document
import os
from src.config import REPORT_DIR, get_logger

logger = get_logger(__name__)

def generate_report(data, filename):
    """ Generate a Word report based on processed data. """
    try:
        doc = Document()
        doc.add_heading('I&W Report', level=1)

        table = doc.add_table(rows=1, cols=5)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Search Term'
        hdr_cells[1].text = 'ASN'
        hdr_cells[2].text = 'Country'
        hdr_cells[3].text = 'Reputation'
        hdr_cells[4].text = 'Link'

        for _, row in data.iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = str(row.get('search_term', 'N/A'))
            row_cells[1].text = str(row.get('base_type', 'N/A'))
            row_cells[2].text = extract_date(row.get('timestamp_vt', 'N/A'))
            row_cells[3].text = str(row.get('observed_by_otx', 'N/A'))
            row_cells[4].text = str(row.get('notes', ''))

        report_path = os.path.join(REPORT_DIR, filename)
        doc.save(report_path)
        logger.info(f"Report saved to {report_path}")

    except Exception as e:
        logger.error(f"Error generating report: {e}")
