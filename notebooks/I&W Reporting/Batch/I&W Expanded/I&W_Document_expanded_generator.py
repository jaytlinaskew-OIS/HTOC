#!/usr/bin/env python3
"""
I&W Document Expanded Generator
Converts the Jupyter notebook functionality into a standalone Python script.
This script generates individual I&W reports for indicators and adds them to reported indicators list.
"""

import sys
import os
import urllib3
import pandas as pd
import requests
import ipaddress
from datetime import datetime
import concurrent.futures
import glob
import docx
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import csv
import warnings

# Suppress warnings
warnings.simplefilter(action='ignore', category=pd.errors.SettingWithCopyWarning)

def setup_paths_and_imports():
    """Setup system paths and import required modules."""
    print("Setting up paths and imports...")
    
    # Add your local ThreatConnect SDK to path
    sys.path.append("Z:/HTOC/Data_Analytics/threatconnect")
    from ThreatConnect import ThreatConnect
    from RequestObject import RequestObject

    # Add your project repo to path
    project_root = r"C:\Users\jaskew\Documents\project_repository\scripts\Data Movement\ThrearConnect-api-pull"
    if project_root not in sys.path:
        sys.path.append(project_root)

    from utils.config_loader import load_config
    
    return ThreatConnect, RequestObject, load_config, project_root

def initialize_threatconnect(project_root, load_config):
    """Initialize ThreatConnect API connection."""
    print("Initializing ThreatConnect connection...")
    
    # Load API config
    config_path = os.path.join(project_root, "utils", "config.json")
    try:
        api_secret_key, api_access_id, api_base_url, api_default_org = load_config(config_path)
        print(f"Loaded config from: {config_path}")
        print(f"Base URL: {api_base_url}")
        print(f"Access ID: {api_access_id}")
        print(f"Default Org: {api_default_org}")
    except Exception as e:
        print(f"[ERROR] Failed to load configuration: {e}")
        sys.exit(1)

    # Disable SSL verification warnings
    urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

    # Initialize ThreatConnect session
    try:
        tc = ThreatConnect(api_access_id, api_secret_key, api_default_org, api_base_url)
        print("ThreatConnect initialized.")
    except Exception as e:
        print(f"[ERROR] Failed to initialize ThreatConnect: {e}")
        sys.exit(1)

    # Define the owner (organization scope)
    owner = 'HTOC Org'

    # Create a request object
    try:
        ro = RequestObject()
        ro.set_http_method('GET')
        ro.set_owner(owner)
        ro.set_owner_allowed(True)
        print("RequestObject successfully created.")
    except Exception as e:
        print(f"[ERROR] Failed to initialize RequestObject: {e}")
        sys.exit(1)
    
    return tc, ro

def load_most_recent_spreadsheet():
    """Load the most recent spreadsheet from the expanded directory."""
    print("Loading most recent spreadsheet...")
    
    spreadsheet_dir = r"Z:\HTOC\HTOC Reports\I&W Reports\5. I&W Staging\Spreadsheet\Expanded"
    
    if not os.path.exists(spreadsheet_dir):
        print(f"Directory not found: {spreadsheet_dir}")
        return None

    # Get all files in the directory
    files = glob.glob(os.path.join(spreadsheet_dir, "*"))
    
    if files:
        most_recent_file = max(files, key=os.path.getmtime)
        print(f"Most recent file: {most_recent_file}")
        
        try:
            df = pd.read_excel(most_recent_file)
            print(f"Loaded {len(df)} rows from spreadsheet")
            return df
        except Exception as e:
            print(f"Error reading file {most_recent_file}: {e}")
            return None
    else:
        print("No files found in the directory.")
        return None

def is_ip(value):
    """Check if the given value is a valid IP address."""
    try:
        ipaddress.ip_address(value)
        return True
    except ValueError:
        return False

def determine_query_type(query):
    """Determine if the query is an IP, domain, or hostname."""
    if is_ip(query):
        return "ip"
    elif "." in query:
        return "hostname"
    else:
        return "domain"

def fetch_virustotal_data(query, tc, ro):
    """Fetch data from VirusTotal API using ThreatConnect enrich endpoint."""
    import urllib.parse

    indicator_values = [query] if isinstance(query, str) else query
    enriched_results = []

    for value in indicator_values:
        try:
            encoded_value = urllib.parse.quote(value)
            enrich_url = f'/v3/indicators/{encoded_value}/enrich?type=Shodan&type=VirusTotalV3'
            ro.set_http_method('POST')
            ro.set_request_uri(enrich_url)
            ro.set_body({})
            enrich_response = tc.api_request(ro)

            if enrich_response.status_code == 200:
                enrich_data = enrich_response.json()
                enrich_data['summary'] = value
                enriched_results.append(enrich_data)
        except Exception as e:
            continue

    if enriched_results:
        flattened = []
        for item in enriched_results:
            if isinstance(item, dict) and 'data' in item:
                flat = item.copy()
                flat.update(pd.json_normalize(item['data']).iloc[0].to_dict() if isinstance(item['data'], dict) else {})
                del flat['data']
                flattened.append(flat)
            else:
                flattened.append(item)
        enriched_results = flattened

    if len(enriched_results) == 1:
        return enriched_results[0]
    
    return enriched_results

def fetch_otx_data(query):
    """Fetch data from OTX API for IP, Domain, or Hostname."""
    # API Keys
    OTX_API_KEY = "ea83cf4792fc5db7acc741e82286c0b717fc99be98ec5b14de7e63cd3f74bcfe"
    OTX_HEADERS = {"X-OTX-API-KEY": OTX_API_KEY}
    
    # API URLs
    OTX_URL_IP = "https://otx.alienvault.io/api/v1/indicators/IPv4/{}/general"
    OTX_URL_DOMAIN = "https://otx.alienvault.io/api/v1/indicators/domain/{}/general"
    OTX_URL_HOSTNAME = "https://otx.alienvault.io/api/v1/indicators/hostname/{}"
    
    query_type = determine_query_type(query)

    if query_type == "ip":
        url = OTX_URL_IP.format(query)
    elif query_type == "hostname":
        url = OTX_URL_HOSTNAME.format(query)
    else:
        url = OTX_URL_DOMAIN.format(query)

    try:
        response = requests.get(url, headers=OTX_HEADERS, verify=False)
        response.raise_for_status()
        data = response.json()

        return {
            "search_term": query,
            "base_indicator": data.get('base_indicator', {}),
            "reputation": data.get('reputation'),
            "geo": data.get('geo', {}),
            "whois": data.get('whois', {}),
            "open_ports": data.get('open_ports', []),
            "link": f"https://otx.alienvault.com/indicator/{query_type}/{query}"
        }

    except Exception as e:
        print(f"OTX Error for {query}: {e}")
        return None

def unnest_base_indicator(df):
    """Unnest the 'base_indicator' column and create separate columns for its keys."""
    if 'base_indicator' not in df.columns:
        return df

    base_df = pd.json_normalize(df['base_indicator'])
    base_df.columns = [f"base_{col}" for col in base_df.columns]

    df = df.drop(columns=['base_indicator']).reset_index(drop=True)
    df = pd.concat([df, base_df], axis=1)
    return df

def process_indicator(indicator, observed_by, partner_count, tc, ro):
    """Fetch data for a single indicator."""
    timestamp = datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S")

    with concurrent.futures.ThreadPoolExecutor() as executor:
        vt_future = executor.submit(fetch_virustotal_data, indicator, tc, ro)
        otx_future = executor.submit(fetch_otx_data, indicator)

        vt_data = vt_future.result()
        otx_data = otx_future.result()

    if vt_data:
        vt_data.update({
            "timestamp": timestamp,
            "observed_by": observed_by,
            "partner_count": partner_count
        })

    if otx_data:
        otx_data.update({
            "timestamp": timestamp,
            "observed_by": observed_by,
            "partner_count": partner_count
        })

    return vt_data, otx_data

def fetch_enrichment_data(recent_tags, tc, ro):
    """Main function to process indicators and fetch enrichment data."""
    print("Fetching enrichment data from VirusTotal and OTX...")
    
    if 'summary' not in recent_tags.columns:
        print("The 'summary' column is missing.")
        return pd.DataFrame(), pd.DataFrame()

    search_terms = recent_tags['summary'].dropna().unique().tolist()
    print(f"Processing {len(search_terms)} unique search terms.")

    vt_results = []
    otx_results = []

    with concurrent.futures.ThreadPoolExecutor() as executor:
        futures = []
        for indicator in search_terms:
            partners = recent_tags.loc[recent_tags['summary'] == indicator, 'partners'].values
            partner_count = recent_tags.loc[recent_tags['summary'] == indicator, 'partner_count'].values
            observed_by = partners[0] if len(partners) > 0 else "N/A"
            partner_count = partner_count[0] if len(partner_count) > 0 else "N/A"

            futures.append(executor.submit(process_indicator, indicator, observed_by, partner_count, tc, ro))

        for future in concurrent.futures.as_completed(futures):
            vt_data, otx_data = future.result()
            if vt_data:
                vt_results.append(vt_data)
            if otx_data:
                otx_results.append(otx_data)

    vt_df = pd.DataFrame(vt_results)
    otx_df = pd.DataFrame(otx_results)
    
    otx_df = unnest_base_indicator(otx_df)

    return vt_df, otx_df

def consolidate_sources(vt_df, otx_df):
    """Consolidate links from both DataFrames for each search term."""
    if vt_df is None or vt_df.empty:
        vt_links = pd.DataFrame(columns=['search_term', 'vt_links'])
    else:
        vt_links = vt_df.groupby('search_term')['link'].apply(lambda x: ', '.join(x.dropna())).reset_index()
        vt_links.columns = ['search_term', 'vt_links']

    if otx_df is None or otx_df.empty:
        otx_links = pd.DataFrame(columns=['search_term', 'otx_links'])
    else:
        otx_links = otx_df.groupby('search_term')['link'].apply(lambda x: ', '.join(x.dropna())).reset_index()
        otx_links.columns = ['search_term', 'otx_links']

    # Merge the two link sets
    if not vt_links.empty and not otx_links.empty:
        consolidated = pd.merge(vt_links, otx_links, on='search_term', how='outer')
    elif not vt_links.empty:
        consolidated = vt_links.copy()
        consolidated['otx_links'] = ''
    elif not otx_links.empty:
        consolidated = otx_links.copy()
        consolidated['vt_links'] = ''
    else:
        return pd.DataFrame(columns=['search_term', 'sources'])

    # Combine the links, handling NaN values
    consolidated['sources'] = consolidated[['vt_links', 'otx_links']].fillna('').apply(
        lambda x: ', '.join(filter(None, x)), axis=1
    )

    return consolidated[['search_term', 'sources']]

def extract_date(timestamp):
    """Extract only the date from the timestamp."""
    if pd.isna(timestamp) or timestamp == 'N/A':
        return 'N/A'
    
    try:
        if isinstance(timestamp, str):
            timestamp = datetime.strptime(timestamp, "%Y-%m-%d %H:%M:%S")
        return timestamp.strftime("%Y-%m-%d")
    except Exception:
        return 'N/A'

def populate_table(table, data):
    """Populate a Word table with the given data."""
    for index, row in data.iterrows():
        new_row = table.add_row().cells
        new_row[0].text = str(row.get('search_term', 'N/A'))
        new_row[1].text = str(row.get('type', 'N/A'))
        new_row[2].text = extract_date(row.get('observed_date', 'N/A'))
        
        # For the 'observed_by_otx' column, stack values instead of comma-separating
        observed_by_otx = ''
        observed_by_list = []

        # Try to get observed_by from OTX
        if 'observed_by_otx' in row and pd.notna(row['observed_by_otx']):
            observed_by_list.extend([v.strip() for v in str(row['observed_by_otx']).split(',') if v.strip()])
        elif 'observed_by' in row and pd.notna(row['observed_by']):
            observed_by_list.extend([v.strip() for v in str(row['observed_by']).split(',') if v.strip()])

        # Remove duplicates and join with newlines
        if observed_by_list:
            observed_by_otx = '\n'.join(sorted(set(observed_by_list)))
        else:
            observed_by_otx = 'N/A'
        new_row[3].text = str(observed_by_otx)
        new_row[4].text = str(row.get('notes', ''))

def fill_word_template(template_path, output_path, df, group_id, output_dir):
    """Fill the template with data and place sources outside the table."""
    if not os.path.exists(template_path):
        print(f"Template not found: {template_path}")
        return
    
    try:
        doc = Document(template_path)

        # Populate the table
        table = None
        for tbl in doc.tables:
            if "Indicators/Identifiers" in tbl.rows[0].cells[0].text:
                table = tbl
                break

        if table:
            populate_table(table, df)

        # Find and replace placeholders outside the table
        for para in doc.paragraphs:
            if "{{ipaddress}}" in para.text:
                ip_address = str(df['search_term'].iloc[0]) if 'search_term' in df.columns else 'N/A'
                para.text = para.text.replace("{{ipaddress}}", ip_address)
            if "{{asn}}" in para.text:
                asn_value = str(df['asn'].iloc[0]) if 'asn' in df.columns and not df['asn'].isna().all() else 'N/A'
                para.text = para.text.replace("{{asn}}", asn_value)
            if "{{whois}}" in para.text:
                whois_value = str(df['whois'].iloc[0]) if 'whois' in df.columns and not df['whois'].isna().all() else 'N/A'
                para.text = para.text.replace("{{whois}}", whois_value)
            if "{{partners}}" in para.text:
                partners_value = ''
                if 'search_term' in df.columns and not df.empty:
                    search_term = df['search_term'].iloc[0]
                    partners_row = df[df['search_term'] == search_term]
                    if not partners_row.empty and 'partners' in partners_row.columns:
                        partners_value = str(partners_row['partners'].iloc[0])
                if not partners_value:
                    partners_value = 'N/A'
                para.text = para.text.replace("{{partners}}", partners_value)
            if "{{weblink}}" in para.text:
                weblink_value = ''
                if 'search_term' in df.columns and not df.empty:
                    search_term = df['search_term'].iloc[0]
                    if 'webLink' in df.columns:
                        match = df[df['search_term'] == search_term]
                        if not match.empty and pd.notna(match['webLink'].iloc[0]):
                            weblink_value = str(match['webLink'].iloc[0])
                    if not weblink_value and 'link' in df.columns:
                        match = df[df['search_term'] == search_term]
                        if not match.empty and pd.notna(match['link'].iloc[0]):
                            weblink_value = str(match['link'].iloc[0])
                para.text = para.text.replace("{{weblink}}", "")
                if weblink_value:
                    r_id = doc.part.relate_to(
                        weblink_value, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True
                    )
                    hyperlink = OxmlElement('w:hyperlink')
                    hyperlink.set(qn('r:id'), r_id)
                    new_run = OxmlElement('w:r')
                    rPr = OxmlElement('w:rPr')
                    rStyle = OxmlElement('w:rStyle')
                    rStyle.set(qn('w:val'), 'Hyperlink')
                    rPr.append(rStyle)
                    new_run.append(rPr)
                    t = OxmlElement('w:t')
                    t.text = weblink_value
                    new_run.append(t)
                    hyperlink.append(new_run)
                    para._p.append(hyperlink)
                else:
                    para.text = "N/A"
            if "{{sources}}" in para.text:
                def add_hyperlink(paragraph, url):
                    part = paragraph.part
                    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
                    hyperlink = OxmlElement('w:hyperlink')
                    hyperlink.set(qn('r:id'), r_id)
                    new_run = OxmlElement('w:r')
                    rPr = OxmlElement('w:rPr')
                    rStyle = OxmlElement('w:rStyle')
                    rStyle.set(qn('w:val'), 'Hyperlink')
                    rPr.append(rStyle)
                    new_run.append(rPr)
                    t = OxmlElement('w:t')
                    t.text = url
                    new_run.append(t)
                    hyperlink.append(new_run)
                    paragraph._p.append(hyperlink)

                para.text = para.text.replace("{{sources}}", "")

                # Add each source as a hyperlink, stacked (one per line, no commas)
                sources = []
                for srcs in df['sources'].dropna().unique():
                    for src in [s.strip() for s in srcs.split(',') if s.strip()]:
                        sources.append(src)
                for i, src in enumerate(sources):
                    add_hyperlink(para, src)
                    if i < len(sources) - 1:
                        para.add_run().add_break()

        # Save the document
        current_date = datetime.now().strftime("%Y-%m-%d")
        folder_path = os.path.join(output_dir, current_date)
        os.makedirs(folder_path, exist_ok=True)

        # Use indicator name for single indicators, group_id for multiple indicators
        if len(df) == 1:
            indicator_name = str(df['search_term'].iloc[0])
            sanitized_name = indicator_name.replace(":", "_").replace("/", "_").replace(" ", "_")
            output_path = os.path.join(folder_path, f"I&W_Report_{sanitized_name}.docx")
        else:
            sanitized_group_id = str(group_id).replace(":", "_").replace("/", "_").replace(" ", "_")
            output_path = os.path.join(folder_path, f"I&W_Report_Group_{sanitized_group_id}.docx")
        doc.save(output_path)
        print(f"Saved report: {output_path}")

    except Exception as e:
        print(f"Error while generating report for group {group_id}: {e}")

def generate_reports(vt_df, otx_df, recent_tags):
    """Generate I&W reports from the enriched data."""
    print("Generating I&W reports...")
    
    # File paths
    TEMPLATE_PATH = r"C:\Users\jaskew\Documents\project_repository\notebooks\I&W Reporting\I&W Report Template.docx"
    OUTPUT_DIR = r"Z:\HTOC\HTOC Reports\I&W Reports\5. I&W Staging\Expanded Reports"
    
    # Ensure output directory exists
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    # Normalize VT columns you rely on
    if vt_df is not None and not vt_df.empty:
        vt_df = vt_df.rename(columns={'ip': 'search_term', 'webLink': 'link'})

    # Normalize recent_tags to 'search_term' so we can join on one key
    if isinstance(recent_tags, pd.DataFrame) and not recent_tags.empty:
        recent_norm = recent_tags.rename(columns={'summary': 'search_term'}).copy()
        keep_cols = [c for c in ['search_term','observations','type','partners','observed_date','webLink','group_id'] if c in recent_norm.columns]
        recent_norm = recent_norm[keep_cols]
    else:
        recent_norm = pd.DataFrame(columns=['search_term'])

    # Combine VT and OTX
    if vt_df is None or vt_df.empty:
        combined_df = otx_df.copy()
    else:
        combined_df = pd.merge(vt_df, otx_df, on='search_term', how='outer', suffixes=('_vt', '_otx'))

    # Consolidate sources
    sources_df = consolidate_sources(vt_df, otx_df)
    combined_df = pd.merge(combined_df, sources_df, on='search_term', how='left')

    # Merge recent tags
    combined_df = pd.merge(combined_df, recent_norm, on='search_term', how='left', suffixes=('', '_tag'))
    
    # Ensure we use the indicator type from recent_tags if available
    if 'type_tag' in combined_df.columns:
        combined_df['type'] = combined_df['type_tag']
        combined_df.drop(columns=['type_tag'], inplace=True)

    # Group by group_id and generate one report per group
    if 'group_id' in combined_df.columns:
        print(f"Grouping indicators by group_id...")
        
        # Get unique group_ids, excluding null/NaN values
        unique_groups = combined_df['group_id'].dropna().unique()
        
        print(f"Found {len(unique_groups)} unique groups to process")
        
        for group_id in unique_groups:
            # Get all indicators for this group
            group_df = combined_df[combined_df['group_id'] == group_id]
            
            if not group_df.empty:
                print(f"Processing group {group_id} with {len(group_df)} indicators")
                fill_word_template(TEMPLATE_PATH, None, group_df, group_id, OUTPUT_DIR)
            else:
                print(f"No indicators found for group {group_id}")
        
        # Handle indicators without group_id (create individual reports)
        ungrouped_df = combined_df[combined_df['group_id'].isna()]
        if not ungrouped_df.empty:
            print(f"Processing {len(ungrouped_df)} ungrouped indicators...")
            for indicator in ungrouped_df['search_term'].dropna().unique():
                indicator_df = ungrouped_df[ungrouped_df['search_term'] == indicator]
                sanitized = str(indicator).replace(":", "_").replace("/", "_")
                fill_word_template(TEMPLATE_PATH, None, indicator_df, sanitized, OUTPUT_DIR)
    else:
        print("No group_id column found, generating individual reports...")
        # Fallback to original behavior if no group_id column
        for indicator in combined_df['search_term'].dropna().unique():
            indicator_df = combined_df[combined_df['search_term'] == indicator]
            sanitized = str(indicator).replace(":", "_").replace("/", "_")
            fill_word_template(TEMPLATE_PATH, None, indicator_df, sanitized, OUTPUT_DIR)

def update_reported_indicators(df):
    """Update the reported indicators CSV with new indicators."""
    print("Updating reported indicators list...")
    
    indicators_csv_path = r"Z:\HTOC\HTOC Reports\I&W Reports\5. I&W Staging\Reported Indicators\indicators.csv"
    
    # Ensure directory exists
    os.makedirs(os.path.dirname(indicators_csv_path), exist_ok=True)
    
    # Collect all unique indicators/summaries from df
    indicators_to_add = df['summary'].dropna().unique()
    
    print(f"Adding {len(indicators_to_add)} indicators to reported list")

    with open(indicators_csv_path, mode='a', newline='', encoding='utf-8') as csvfile:
        writer = csv.writer(csvfile)
        for indicator in indicators_to_add:
            writer.writerow([indicator, " - "])

def main():
    """Main execution function."""
    print("=== I&W Document Expanded Generator ===")
    print(f"Started at: {datetime.now()}")
    
    try:
        # Setup and initialization
        ThreatConnect, RequestObject, load_config, project_root = setup_paths_and_imports()
        tc, ro = initialize_threatconnect(project_root, load_config)
        
        # Load most recent spreadsheet
        df = load_most_recent_spreadsheet()
        
        if df is None or df.empty:
            print("No data loaded from spreadsheet. Exiting.")
            return
        
        print(f"Processing {len(df)} indicators from spreadsheet")
        
        # Fetch enrichment data
        vt_df, otx_df = fetch_enrichment_data(df, tc, ro)
        
        # Generate reports
        if vt_df is None or vt_df.empty:
            generate_reports(None, otx_df, df)
        else:
            generate_reports(vt_df, otx_df, df)
        
        # Update reported indicators
        update_reported_indicators(df)
        
        print("=== Processing Complete ===")
        print(f"Generated reports for indicators from spreadsheet")
        print(f"Updated reported indicators list")
        print(f"Completed at: {datetime.now()}")
        
    except Exception as e:
        print(f"[ERROR] Script failed: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)

if __name__ == "__main__":
    main()