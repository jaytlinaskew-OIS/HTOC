#!/usr/bin/env python3
"""
I&W Report Generator Script

This script processes indicators from Excel files and generates Word reports
with threat intelligence data from VirusTotal and OTX APIs.

Converted from: I&W_Spreadsheet_Generate_Report.ipynb
Author: GitHub Copilot
Created: October 8, 2025
"""

import sys
import os
import pandas as pd
import requests
import ipaddress
from datetime import datetime, timezone
import concurrent.futures
import urllib.parse
from docx import Document


def main():
    """Main execution function."""
    print("Starting I&W Report Generator...")
    
    # ═══════════════════════════════════════════════════════════════════════════════
    # CONFIGURATION
    # ═══════════════════════════════════════════════════════════════════════════════
    
    # API Keys
    OTX_API_KEY = "ea83cf4792fc5db7acc741e82286c0b717fc99be98ec5b14de7e63cd3f74bcfe"

    # File paths
    TEMPLATE_PATH = r"z:\HTOC\HTOC Reports\I&W Reports\5. I&W Staging\I&W Report Template.docx"
    OUTPUT_DIR = r"z:\HTOC\HTOC Reports\I&W Reports\5. I&W Staging\Generated Reports"

    # Input Excel file path - look for the most recent file
    excel_dir = r"Z:\HTOC\HTOC Reports\I&W Reports\5. I&W Staging\Spreadsheet"
    excel_pattern = "I&W_indicators_full_*.xlsx"
    
    # Find the most recent Excel file
    excel_file_path = find_latest_excel_file(excel_dir, excel_pattern)
    if not excel_file_path:
        print("ERROR: No I&W indicators Excel file found in the specified directory")
        sys.exit(1)
    
    print(f"Using Excel file: {excel_file_path}")
    
    # Ensure output directory exists
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    # ═══════════════════════════════════════════════════════════════════════════════
    # DATA LOADING
    # ═══════════════════════════════════════════════════════════════════════════════
    
    print("Loading indicators from Excel file...")
    try:
        filtered_recent_tags = pd.read_excel(excel_file_path)
        print(f"Loaded {len(filtered_recent_tags)} indicators from Excel file")
    except Exception as e:
        print(f"ERROR: Failed to load Excel file: {e}")
        sys.exit(1)

    # ═══════════════════════════════════════════════════════════════════════════════
    # THREAT INTELLIGENCE COLLECTION
    # ═══════════════════════════════════════════════════════════════════════════════
    
    print("Collecting threat intelligence from APIs...")
    
    # Initialize ThreatConnect if needed for VirusTotal enrichment
    try:
        # Add ThreatConnect SDK to path
        sys.path.append("Z:/HTOC/Data_Analytics/threatconnect")
        from ThreatConnect import ThreatConnect
        from RequestObject import RequestObject
        
        # Load API config - using the same approach as I&W Spreadsheet
        project_root = r"Z:\HTOC\HTOC Reports\I&W Reports\5. I&W Staging\I&W Report Processing Scripts"
        if project_root not in sys.path:
            sys.path.append(project_root)

        # Add the scripts directory to the path to import config_loader
        scripts_path = os.path.join(project_root, "scripts")
        if scripts_path not in sys.path:
            sys.path.append(scripts_path)

        from config_loader import get_threatconnect_config

        config_path = os.path.join(project_root, "utils", "config.json")
        tc_config = get_threatconnect_config(config_path)
        api_secret_key = tc_config["secret_key"]
        api_access_id = tc_config["access_id"]
        api_base_url = tc_config["base_url"]
        api_default_org = tc_config["default_org"]
        
        # Initialize ThreatConnect
        global tc, ro
        tc = ThreatConnect(api_access_id, api_secret_key, api_default_org, api_base_url)
        ro = RequestObject()
        print("ThreatConnect initialized for VirusTotal enrichment")
        
    except Exception as e:
        print(f"Warning: Could not initialize ThreatConnect: {e}")
        print("Will skip VirusTotal enrichment")
        tc = None
        ro = None

    # Collect threat intelligence
    vt_df, otx_df = collect_threat_intelligence(filtered_recent_tags, OTX_API_KEY)

    # ═══════════════════════════════════════════════════════════════════════════════
    # REPORT GENERATION
    # ═══════════════════════════════════════════════════════════════════════════════
    
    print("Generating Word reports...")
    
    # Generate reports
    generate_reports(vt_df, otx_df, filtered_recent_tags, TEMPLATE_PATH, OUTPUT_DIR)
    
    print("I&W Report Generator completed successfully!")


def find_latest_excel_file(directory, pattern):
    """Find the most recent Excel file matching the pattern."""
    import glob
    
    search_path = os.path.join(directory, pattern)
    files = glob.glob(search_path)
    
    if not files:
        return None
    
    # Sort by modification time, newest first
    files.sort(key=os.path.getmtime, reverse=True)
    return files[0]


def collect_threat_intelligence(recent_tags, otx_api_key):
    """Collect threat intelligence from APIs."""
    
    # Headers for OTX API requests
    OTX_HEADERS = {"X-OTX-API-KEY": otx_api_key}

    if 'summary' not in recent_tags.columns:
        print("The 'summary' column is missing.")
        return pd.DataFrame(), pd.DataFrame()

    search_terms = recent_tags['summary'].dropna().unique().tolist()
    print(f"Processing {len(search_terms)} unique search terms.")

    vt_results = []
    otx_results = []

    with concurrent.futures.ThreadPoolExecutor() as executor:
        futures = []
        for indicator in search_terms:
            partners = recent_tags.loc[recent_tags['summary'] == indicator, 'partners'].values
            partner_count = recent_tags.loc[recent_tags['summary'] == indicator, 'partner_count'].values
            observed_by = partners[0] if len(partners) > 0 else "N/A"
            partner_count = partner_count[0] if len(partner_count) > 0 else "N/A"

            futures.append(executor.submit(process_indicator, indicator, observed_by, partner_count, OTX_HEADERS))

        for future in concurrent.futures.as_completed(futures):
            vt_data, otx_data = future.result()
            if vt_data:
                vt_results.append(vt_data)
            if otx_data:
                otx_results.append(otx_data)

    vt_df = pd.DataFrame(vt_results)
    otx_df = pd.DataFrame(otx_results)
    
    # Unnest base_indicator data
    otx_df = unnest_base_indicator(otx_df)

    return vt_df, otx_df


def is_ip(value):
    """Check if the given value is a valid IP address."""
    try:
        ipaddress.ip_address(value)
        return True
    except ValueError:
        return False


def determine_query_type(query):
    """Determine if the query is an IP, domain, or hostname."""
    if is_ip(query):
        return "ip"
    elif "." in query:
        return "hostname"
    else:
        return "domain"


def fetch_virustotal_data(query):
    """Fetch data from VirusTotal API using ThreatConnect enrich endpoint."""
    global tc, ro
    
    if tc is None or ro is None:
        return None
    
    try:
        encoded_value = urllib.parse.quote(query)
        enrich_url = f'/v3/indicators/{encoded_value}/enrich?type=Shodan&type=VirusTotalV3'
        ro.set_http_method('POST')
        ro.set_request_uri(enrich_url)
        ro.set_body({})
        enrich_response = tc.api_request(ro)

        if enrich_response.status_code == 200:
            enrich_data = enrich_response.json()
            enrich_data['summary'] = query
            
            # Flatten nested data
            if 'data' in enrich_data and isinstance(enrich_data['data'], dict):
                flat_data = enrich_data.copy()
                flat_data.update(pd.json_normalize(enrich_data['data']).iloc[0].to_dict())
                del flat_data['data']
                return flat_data
            
            return enrich_data
            
    except Exception as e:
        print(f"VirusTotal Error for {query}: {e}")
    
    return None


def fetch_otx_data(query, otx_headers):
    """Fetch data from OTX API for IP, Domain, or Hostname."""
    query_type = determine_query_type(query)

    # API URLs
    OTX_URL_IP = "https://otx.alienvault.io/api/v1/indicators/IPv4/{}/general"
    OTX_URL_DOMAIN = "https://otx.alienvault.io/api/v1/indicators/domain/{}/general"
    OTX_URL_HOSTNAME = "https://otx.alienvault.io/api/v1/indicators/hostname/{}"

    if query_type == "ip":
        url = OTX_URL_IP.format(query)
    elif query_type == "hostname":
        url = OTX_URL_HOSTNAME.format(query)
    else:
        url = OTX_URL_DOMAIN.format(query)

    try:
        response = requests.get(url, headers=otx_headers, verify=False)
        response.raise_for_status()
        data = response.json()

        return {
            "search_term": query,
            "base_indicator": data.get('base_indicator', {}),
            "reputation": data.get('reputation'),
            "geo": data.get('geo', {}),
            "whois": data.get('whois', {}),
            "open_ports": data.get('open_ports', []),
            "link": f"https://otx.alienvault.com/indicator/{query_type}/{query}"
        }

    except Exception as e:
        print(f"OTX Error for {query}: {e}")
        return None


def unnest_base_indicator(df):
    """Unnest the 'base_indicator' column and create separate columns for its keys."""
    if df.empty or 'base_indicator' not in df.columns:
        return df

    base_df = pd.json_normalize(df['base_indicator'])
    base_df.columns = [f"base_{col}" for col in base_df.columns]

    df = df.drop(columns=['base_indicator']).reset_index(drop=True)
    df = pd.concat([df, base_df], axis=1)
    return df


def process_indicator(indicator, observed_by, partner_count, otx_headers):
    """Fetch data for a single indicator."""
    timestamp = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S")

    # Fetch data from both APIs
    vt_data = fetch_virustotal_data(indicator)
    otx_data = fetch_otx_data(indicator, otx_headers)

    if vt_data:
        vt_data.update({
            "timestamp": timestamp,
            "observed_by": observed_by,
            "partner_count": partner_count
        })

    if otx_data:
        otx_data.update({
            "timestamp": timestamp,
            "observed_by": observed_by,
            "partner_count": partner_count
        })

    return vt_data, otx_data


def consolidate_sources(vt_df, otx_df):
    """Consolidate links from both DataFrames for each search term."""
    # Handle None or empty DataFrames
    if vt_df is None or vt_df.empty:
        vt_links = pd.DataFrame(columns=['search_term', 'vt_links'])
    else:
        vt_links = vt_df.groupby('search_term')['link'].apply(lambda x: ', '.join(x.dropna())).reset_index()
        vt_links.columns = ['search_term', 'vt_links']

    if otx_df is None or otx_df.empty:
        otx_links = pd.DataFrame(columns=['search_term', 'otx_links'])
    else:
        otx_links = otx_df.groupby('search_term')['link'].apply(lambda x: ', '.join(x.dropna())).reset_index()
        otx_links.columns = ['search_term', 'otx_links']

    # Merge the two link sets
    if not vt_links.empty and not otx_links.empty:
        consolidated = pd.merge(vt_links, otx_links, on='search_term', how='outer')
    elif not vt_links.empty:
        consolidated = vt_links.copy()
        consolidated['otx_links'] = ''
    elif not otx_links.empty:
        consolidated = otx_links.copy()
        consolidated['vt_links'] = ''
    else:
        return pd.DataFrame(columns=['search_term', 'sources'])

    # Combine the links, handling NaN values
    consolidated['sources'] = consolidated[['vt_links', 'otx_links']].fillna('').apply(
        lambda x: ', '.join(filter(None, x)), axis=1
    )

    return consolidated[['search_term', 'sources']]


def extract_date(timestamp):
    """Extract only the date from the timestamp."""
    if pd.isna(timestamp) or timestamp == 'N/A':
        return 'N/A'
    
    # Handle datetime object or string
    try:
        # Attempt to parse as a datetime object
        if isinstance(timestamp, str):
            timestamp = datetime.strptime(timestamp, "%Y-%m-%d %H:%M:%S")
        return timestamp.strftime("%Y-%m-%d")
    except Exception:
        return 'N/A'


def populate_table(table, data):
    """Populate a Word table with the given data."""
    # Iterate over data and populate rows
    for index, row in data.iterrows():
        # Insert a new row
        new_row = table.add_row().cells
        new_row[0].text = str(row.get('search_term', 'N/A'))
        new_row[1].text = str(row.get('type', 'N/A'))
        new_row[2].text = extract_date(row.get('observed_date', 'N/A'))
        
        # Handle observed_by column
        observed_by_otx = ''
        observed_by_list = []

        # Try to get observed_by from OTX
        if 'observed_by_otx' in row and pd.notna(row['observed_by_otx']):
            observed_by_list.extend([v.strip() for v in str(row['observed_by_otx']).split(',') if v.strip()])
        elif 'observed_by' in row and pd.notna(row['observed_by']):
            observed_by_list.extend([v.strip() for v in str(row['observed_by']).split(',') if v.strip()])

        # Remove duplicates and join with newlines
        if observed_by_list:
            observed_by_otx = '\n'.join(sorted(set(observed_by_list)))
        else:
            observed_by_otx = 'N/A'
        new_row[3].text = str(observed_by_otx)
        new_row[4].text = str(row.get('notes', ''))


def fill_word_template(template_path, df, group_id, output_dir):
    """Fill the template with data and place sources outside the table."""
    if not os.path.exists(template_path):
        print(f"Template not found: {template_path}")
        return
    
    try:
        doc = Document(template_path)

        # Populate the table
        table = None
        for tbl in doc.tables:
            if len(tbl.rows) > 0 and "Indicators/Identifiers" in tbl.rows[0].cells[0].text:
                table = tbl
                break

        if table:
            populate_table(table, df)

        # Find and replace placeholders outside the table
        for para in doc.paragraphs:
            if "{{ipaddress}}" in para.text:
                ip_address = str(df['search_term'].iloc[0]) if 'search_term' in df.columns else 'N/A'
                para.text = para.text.replace("{{ipaddress}}", ip_address)
            if "{{asn}}" in para.text:
                asn_value = str(df['asn'].iloc[0]) if 'asn' in df.columns and not df['asn'].isna().all() else 'N/A'
                para.text = para.text.replace("{{asn}}", asn_value)
            if "{{whois}}" in para.text:
                whois_value = str(df['whois'].iloc[0]) if 'whois' in df.columns and not df['whois'].isna().all() else 'N/A'
                para.text = para.text.replace("{{whois}}", whois_value)
            if "{{partners}}" in para.text:
                partners_value = str(df['partners'].iloc[0]) if 'partners' in df.columns and not df['partners'].isna().all() else 'N/A'
                para.text = para.text.replace("{{partners}}", partners_value)
            if "{{weblink}}" in para.text:
                weblink_value = ''
                if 'search_term' in df.columns and not df.empty:
                    search_term = df['search_term'].iloc[0]
                    # Try to find a 'webLink' in df for the indicator
                    if 'webLink' in df.columns:
                        match = df[df['search_term'] == search_term]
                        if not match.empty and pd.notna(match['webLink'].iloc[0]):
                            weblink_value = str(match['webLink'].iloc[0])
                    # Fallback: try 'link' column
                    if not weblink_value and 'link' in df.columns:
                        match = df[df['search_term'] == search_term]
                        if not match.empty and pd.notna(match['link'].iloc[0]):
                            weblink_value = str(match['link'].iloc[0])
                
                para.text = para.text.replace("{{weblink}}", weblink_value if weblink_value else "N/A")
                
            if "{{sources}}" in para.text:
                # Replace with sources as plain text (simplified)
                sources_text = ''
                if 'sources' in df.columns and not df['sources'].isna().all():
                    sources_text = str(df['sources'].iloc[0])
                para.text = para.text.replace("{{sources}}", sources_text if sources_text else "N/A")

        # Save the document
        current_date = datetime.now().strftime("%Y-%m-%d")
        folder_path = os.path.join(output_dir, current_date)
        os.makedirs(folder_path, exist_ok=True)

        # Use indicator name for single indicators, group_id for multiple indicators
        if len(df) == 1:
            indicator_name = str(df['search_term'].iloc[0])
            sanitized_name = indicator_name.replace(":", "_").replace("/", "_").replace(" ", "_")
            output_path = os.path.join(folder_path, f"I&W_Report_{sanitized_name}.docx")
        else:
            sanitized_group_id = str(group_id).replace(":", "_").replace("/", "_").replace(" ", "_")
            output_path = os.path.join(folder_path, f"I&W_Report_Group_{sanitized_group_id}.docx")
        
        doc.save(output_path)
        print(f"Saved report: {output_path}")

    except Exception as e:
        print(f"Error while generating report for group {group_id}: {e}")


def generate_reports(vt_df, otx_df, recent_tags, template_path, output_dir):
    """Generate Word reports from the collected data."""
    
    # Normalize VT columns
    if vt_df is not None and not vt_df.empty:
        vt_df = vt_df.rename(columns={'ip': 'search_term', 'webLink': 'link'})

    # Normalize recent_tags to 'search_term'
    if isinstance(recent_tags, pd.DataFrame) and not recent_tags.empty:
        recent_norm = recent_tags.rename(columns={'summary': 'search_term'}).copy()
        keep_cols = [c for c in ['search_term','observations','type','partners','observed_date','webLink','group_id'] if c in recent_norm.columns]
        recent_norm = recent_norm[keep_cols]
    else:
        recent_norm = pd.DataFrame(columns=['search_term'])

    # Combine VT and OTX data
    if vt_df is None or vt_df.empty:
        combined_df = otx_df.copy() if otx_df is not None and not otx_df.empty else pd.DataFrame()
    else:
        if otx_df is not None and not otx_df.empty:
            combined_df = pd.merge(vt_df, otx_df, on='search_term', how='outer', suffixes=('_vt', '_otx'))
        else:
            combined_df = vt_df.copy()

    if combined_df.empty:
        print("No threat intelligence data collected. Cannot generate reports.")
        return

    # Consolidate sources
    sources_df = consolidate_sources(vt_df, otx_df)
    if not sources_df.empty:
        combined_df = pd.merge(combined_df, sources_df, on='search_term', how='left')

    # Merge recent tags
    combined_df = pd.merge(combined_df, recent_norm, on='search_term', how='left', suffixes=('', '_tag'))
    
    # Use the indicator type from recent_tags if available
    if 'type_tag' in combined_df.columns:
        combined_df['type'] = combined_df['type_tag']
        combined_df.drop(columns=['type_tag'], inplace=True)

    # Generate reports by group_id if available
    if 'group_id' in combined_df.columns:
        print("Grouping indicators by group_id...")
        
        unique_groups = combined_df['group_id'].dropna().unique()
        print(f"Found {len(unique_groups)} unique groups to process")
        
        for group_id in unique_groups:
            group_df = combined_df[combined_df['group_id'] == group_id]
            
            if not group_df.empty:
                print(f"Processing group {group_id} with {len(group_df)} indicators")
                fill_word_template(template_path, group_df, group_id, output_dir)

        # Handle indicators without group_id
        ungrouped_df = combined_df[combined_df['group_id'].isna()]
        if not ungrouped_df.empty:
            print(f"Processing {len(ungrouped_df)} ungrouped indicators...")
            for indicator in ungrouped_df['search_term'].dropna().unique():
                indicator_df = ungrouped_df[ungrouped_df['search_term'] == indicator]
                sanitized = str(indicator).replace(":", "_").replace("/", "_")
                fill_word_template(template_path, indicator_df, sanitized, output_dir)
    else:
        print("No group_id column found, generating individual reports...")
        for indicator in combined_df['search_term'].dropna().unique():
            indicator_df = combined_df[combined_df['search_term'] == indicator]
            sanitized = str(indicator).replace(":", "_").replace("/", "_")
            fill_word_template(template_path, indicator_df, sanitized, output_dir)


if __name__ == "__main__":
    main()