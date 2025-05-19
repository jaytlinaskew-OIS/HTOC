import os
import pandas as pd
import docx
from datetime import datetime
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement

# File paths
TEMPLATE_PATH = r"z:\\HTOC\\HTOC Reports\\I&W Reports\\5. I&W Staging\\I&W Report Template.docx"
OUTPUT_DIR = r"z:\\HTOC\\HTOC Reports\\I&W Reports\\5. I&W Staging\\Generated Reports"

# Ensure output directory exists
os.makedirs(OUTPUT_DIR, exist_ok=True)


def consolidate_sources(vt_df, otx_df):
    """ Consolidate links from both DataFrames for each search term. """
    # Collect links from VirusTotal
    vt_links = vt_df.groupby('search_term')['link'].apply(lambda x: ', '.join(x.dropna())).reset_index()
    vt_links.columns = ['search_term', 'vt_links']

    # Collect links from OTX
    otx_links = otx_df.groupby('search_term')['link'].apply(lambda x: ', '.join(x.dropna())).reset_index()
    otx_links.columns = ['search_term', 'otx_links']

    # Merge the two link sets
    consolidated = pd.merge(vt_links, otx_links, on='search_term', how='outer')

    # Combine the links, handling NaN values
    consolidated['sources'] = consolidated[['vt_links', 'otx_links']].fillna('').apply(
        lambda x: ', '.join(filter(None, x)), axis=1
    )

    return consolidated[['search_term', 'sources']]

def extract_date(timestamp):
    """ Extract only the date from the timestamp. """
    if pd.isna(timestamp) or timestamp == 'N/A':
        return 'N/A'
    
    # Handle datetime object or string
    try:
        # Attempt to parse as a datetime object
        if isinstance(timestamp, str):
            timestamp = datetime.strptime(timestamp, "%Y-%m-%d %H:%M:%S")
        return timestamp.strftime("%Y-%m-%d")
    except Exception:
        return 'N/A'

def populate_table(table, data):
    """ Populate a Word table with the given data. """
    # Iterate over data and populate rows
    for index, row in data.iterrows():
        # Insert a new row before the last row (template row)
        new_row = table.add_row().cells
        new_row[0].text = str(row.get('search_term', 'N/A'))
        new_row[1].text = str(row.get('type', 'N/A'))
        new_row[2].text = extract_date(row.get('observed_date', 'N/A'))
        # For the 'observed_by_otx' column, stack values instead of comma-separating
        observed_by_otx = row.get('observed_by_otx', 'N/A')
        if pd.notna(observed_by_otx) and isinstance(observed_by_otx, str):
            # Split by comma, strip whitespace, and join with newlines
            observed_by_otx = '\n'.join([v.strip() for v in observed_by_otx.split(',') if v.strip()])
        new_row[3].text = str(observed_by_otx)
        new_row[4].text = str(row.get('notes', ''))
        
def fill_word_template(template_path, output_path, df):
    """ Fill the template with data and place sources outside the table. """
    if not os.path.exists(template_path):
        print(f"Template not found: {template_path}")
        return
    
    try:
        doc = Document(template_path)

        # Populate the table
        table = None
        for tbl in doc.tables:
            if "Indicators/Identifiers" in tbl.rows[0].cells[0].text:
                table = tbl
                break

        if table:
            populate_table(table, df)

        # Find and replace `{{sources}}` placeholder outside the table
        for para in doc.paragraphs:
            if "{{ipaddress}}" in para.text:
                # Try to get the first search_term as the IP address
                ip_address = str(df['search_term'].iloc[0]) if 'search_term' in df.columns else 'N/A'
                para.text = para.text.replace("{{ipaddress}}", ip_address)
            if "{{asn}}" in para.text:
                # Try to get ASN from vt_df if available, else use 'N/A'
                asn_value = str(df['asn'].iloc[0]) if 'asn' in df.columns and not df.empty else 'N/A'
                para.text = para.text.replace("{{asn}}", asn_value)
            if "{{whois}}" in para.text:
                # Try to get WHOIS info from otx_df if available, else use 'N/A'
                whois_value = str(df['whois'].iloc[0]) if 'whois' in df.columns and not df.empty else 'N/A'
                para.text = para.text.replace("{{whois}}", whois_value)
            if "{{partners}}" in para.text:
                # Try to get partners from recent_tags if available, else use 'N/A'
                partners_value = ''
                if 'search_term' in df.columns and not df.empty:
                    search_term = df['search_term'].iloc[0]
                    partners_row = df[df['summary'] == search_term]
                    if not partners_row.empty and 'partners' in partners_row.columns:
                        partners_value = str(partners_row['partners'].iloc[0])
                if not partners_value:
                    partners_value = 'N/A'
                para.text = para.text.replace("{{partners}}", partners_value)
            if "{{weblink}}" in para.text:
                weblink_value = ''
                if 'search_term' in df.columns and not df.empty:
                    search_term = df['search_term'].iloc[0]
                    match = df[df['indicator'] == search_term]
                    if not match.empty and 'webLink' in match.columns:
                        weblink_value = str(match['webLink'].iloc[0])
                para.text = para.text.replace("{{weblink}}", "")
                if weblink_value:
                    # Add hyperlink using WordprocessingML (only show the link, no display text)
                    r_id = doc.part.relate_to(
                        weblink_value, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True
                    )
                    hyperlink = OxmlElement('w:hyperlink')
                    hyperlink.set(qn('r:id'), r_id)
                    new_run = OxmlElement('w:r')
                    rPr = OxmlElement('w:rPr')
                    rStyle = OxmlElement('w:rStyle')
                    rStyle.set(qn('w:val'), 'Hyperlink')
                    rPr.append(rStyle)
                    new_run.append(rPr)
                    t = OxmlElement('w:t')
                    t.text = weblink_value
                    new_run.append(t)
                    hyperlink.append(new_run)
                    para._p.append(hyperlink)
                else:
                    para.text = "N/A"
            if "{{sources}}" in para.text:
                # Build sources as hyperlinks if possible

                # Helper to add a hyperlink to a paragraph
                def add_hyperlink(paragraph, url):
                    # Create the w:hyperlink tag and add needed values
                    part = paragraph.part
                    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
                    hyperlink = OxmlElement('w:hyperlink')
                    hyperlink.set(qn('r:id'), r_id)
                    # Create a w:r element
                    new_run = OxmlElement('w:r')
                    # Create a w:rPr element
                    rPr = OxmlElement('w:rPr')
                    # Add color and underline for hyperlink style
                    rStyle = OxmlElement('w:rStyle')
                    rStyle.set(qn('w:val'), 'Hyperlink')
                    rPr.append(rStyle)
                    new_run.append(rPr)
                    # Create a w:t element and set the text
                    t = OxmlElement('w:t')
                    t.text = url
                    new_run.append(t)
                    hyperlink.append(new_run)
                    paragraph._p.append(hyperlink)

                # Remove the placeholder
                para.text = para.text.replace("{{sources}}", "")

                # Add each source as a hyperlink, separated by commas
                sources = []
                for srcs in df['sources'].dropna().unique():
                    for src in [s.strip() for s in srcs.split(',') if s.strip()]:
                        sources.append(src)
                for i, src in enumerate(sources):
                    add_hyperlink(para, src)
                    if i < len(sources) - 1:
                        para.add_run(", ")

        # Save the modified document
        doc.save(output_path)
        print(f"Saved report: {output_path}")

    except Exception as e:
        print(f"Error while generating report: {e}")

def generate_report(vt_df, otx_df, processed_data):
    """Main function to handle report generation."""

    # Combine VirusTotal and OTX dataframes
    combined_df = pd.merge(
        vt_df, otx_df, on='search_term', how='outer', suffixes=('_vt', '_otx')
    )

    # Consolidate sources into a single dataframe
    sources_df = consolidate_sources(vt_df, otx_df)
    combined_df = pd.merge(
        combined_df, sources_df, on='search_term', how='left'
    )

    # Clearly aggregate observations to avoid duplication
    if not processed_data.empty:
        # Aggregate observations per indicator to single rows
        agg_processed_data = (
            processed_data.groupby('summary', as_index=False)
            .agg({
                'observed_date': lambda x: ', '.join(sorted(set(x.astype(str)))),
                'partners': lambda x: ', '.join(sorted(set(', '.join(x.dropna()).split(', ')))),
                'type': 'first',
                'observations': 'first'
            })
        )

        combined_df = pd.merge(
            combined_df,
            agg_processed_data,
            left_on='search_term',
            right_on='summary',
            how='left'
        )

        # Remove redundant 'summary' after merge
        combined_df.drop(columns=['summary'], inplace=True)

    # Remove any remaining exact duplicates
    combined_df.drop_duplicates(subset=['search_term', 'type', 'observed_date', 'partners'], inplace=True)

    # Define output file path
    current_date = pd.Timestamp.now().strftime("%Y-%m-%d")
    output_file = os.path.join(OUTPUT_DIR, f"I&W_Report_{current_date}.docx")

    # Fill the template table with the clearly deduplicated data
    fill_word_template(TEMPLATE_PATH, output_file, combined_df)


